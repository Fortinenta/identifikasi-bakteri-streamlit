import streamlit as st
import pandas as pd
import io
import time
import os
import json
from auth import get_authenticated_session
from bacdive_mapper import (
    fetch_and_cache_profiles_by_taxonomy,
    calculate_weighted_similarity,
    normalize_columns,
    get_single_strain_json,
    WEIGHTS 
)

# --- 1. Konfigurasi Aplikasi ---
st.set_page_config(
    page_title="🧬 Identifikasi Bakteri (BacDive API)",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- 2. Inisialisasi Sesi ---
@st.cache_resource
def init_session():
    try:
        email = st.secrets["bacdive"]["email"]
        password = st.secrets["bacdive"]["password"]
        session = get_authenticated_session(email, password)
        if session:
            return session
        else:
            st.error("Autentikasi BacDive gagal. Periksa kredensial Anda di secrets.toml.")
            return None
    except KeyError:
        st.error("Kredensial BacDive (email/password) tidak ditemukan di secrets.toml.")
        return None
    except Exception as e:
        st.error(f"Gagal menginisialisasi sesi: {e}")
        return None

# --- 3. PERBAIKAN: Fungsi Debug untuk Testing ---
def test_json_structure(session, genus="Bacillus"):
    """Fungsi untuk testing struktur JSON response."""
    st.subheader("🔍 Testing Struktur JSON BacDive")
    
    with st.expander("Testing API Response Structure", expanded=True):
        if st.button("Test API Response"):
            with st.spinner(f"Testing API response untuk genus '{genus}'..."):
                result = get_single_strain_json(session, genus)
                
                if "error" in result:
                    st.error(f"Error: {result['error']}")
                    if "details" in result:
                        st.code(result["details"])
                else:
                    st.success(f"✅ API Response berhasil untuk strain ID: {result['bacdive_id']}")
                    
                    # Debug info
                    if "debug_info" in result:
                        debug = result["debug_info"]
                        st.info(f"Total strains found: {debug.get('total_strains_found', 0)}")
                        st.info(f"JSON structure keys: {debug.get('json_structure_keys', [])}")
                    
                    # Show JSON structure
                    data = result["data"]
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.subheader("JSON Structure Keys:")
                        if isinstance(data, dict):
                            for key in data.keys():
                                st.write(f"- {key}")
                                if isinstance(data[key], dict) and len(data[key]) < 10:
                                    for subkey in list(data[key].keys())[:5]:
                                        st.write(f"  - {subkey}")
                    
                    with col2:
                        st.subheader("Sample Data:")
                        
                        # Check taxonomy
                        if "Name and taxonomic classification" in data:
                            taxonomy = data["Name and taxonomic classification"]
                            st.write("**Taxonomy Data:**")
                            st.json(taxonomy)
                        
                        # Check morphology  
                        if "Morphology" in data:
                            morphology = data["Morphology"]
                            st.write("**Morphology Data:**")
                            if isinstance(morphology, dict) and morphology:
                                st.json({k: v for k, v in morphology.items() if k in ['cell morphology', 'motility', 'gram stain']})

# --- 4. PERBAIKAN: Logika Inti dengan Enhanced Logging ---
def process_sample(session, user_input, log_container):
    """Fungsi utama untuk memproses satu sampel: fetch, cache, dan analisis."""
    genus = user_input.get("Genus")
    if not genus or pd.isna(genus):
        st.warning("Kolom 'Genus' tidak ditemukan atau kosong untuk sampel ini. Sampel dilewati.")
        return []

    status_placeholder = st.empty()
    
    # PERBAIKAN: Enhanced logging dan error handling
    log_container.info(f"🔄 Memulai proses untuk genus: {genus}")
    
    try:
        raw_profiles = fetch_and_cache_profiles_by_taxonomy(session, genus, status_placeholder, log_container)
        
        # PERBAIKAN: Debugging informasi yang lebih detail
        log_container.info(f"📊 Raw profiles received: {len(raw_profiles) if raw_profiles else 0}")
        if raw_profiles:
            # Log beberapa nama bakteri yang ditemukan
            sample_names = []
            for i, (bid, profile) in enumerate(raw_profiles.items()):
                if i < 5:  # Hanya 5 pertama
                    sample_names.append(f"ID {bid}: {profile.get('Nama Bakteri', 'N/A')}")
            log_container.info(f"📋 Sample bacteria found: {sample_names}")
        
    except Exception as e:
        log_container.error(f"❌ Error dalam fetch_and_cache_profiles_by_taxonomy: {str(e)}")
        log_container.exception(e)
        status_placeholder.error(f"Error saat mengambil profil: {e}")
        return []
    
    if not raw_profiles:
        status_placeholder.warning(f"Tidak ada profil yang ditemukan untuk genus '{genus}'.")
        log_container.warning(f"⚠️ No profiles returned for genus: {genus}")
        time.sleep(3)
        status_placeholder.empty()
        return []
    
    status_placeholder.success(f"Ditemukan {len(raw_profiles)} profil untuk genus '{genus}'. Memulai analisis perbandingan...")
    log_container.info(f"✅ Found {len(raw_profiles)} valid profiles. Starting similarity analysis...")
    time.sleep(2)

    # PERBAIKAN: Melanjutkan bagian yang terpotong
    identification_results = []
    progress_bar = st.progress(0)
    total_profiles = len(raw_profiles)

    for i, (bacdive_id, bacdive_profile) in enumerate(raw_profiles.items()):
        status_placeholder.text(f"⚙️ Membandingkan dengan profil {i+1} dari {total_profiles} (ID: {bacdive_id})...")
        
        try:
            score, details = calculate_weighted_similarity(user_input, bacdive_profile)
            
            # PERBAIKAN: Enhanced logging untuk debugging similarity calculation
            log_container.info(f"🧮 Similarity calculated for ID {bacdive_id}: {score:.2f}%")
            
            if score > 0:
                identification_results.append({
                    "Rank": 0,
                    "Nama Bakteri": bacdive_profile.get("Nama Bakteri", "N/A"),
                    "Persentase": score,
                    "ID": bacdive_id,
                    "details": details
                })
        except Exception as e:
            log_container.error(f"❌ Error calculating similarity for ID {bacdive_id}: {str(e)}")
            continue
            
        progress_bar.progress((i + 1) / total_profiles)

    status_placeholder.text("✅ Perbandingan selesai!")
    time.sleep(1)
    status_placeholder.empty()
    progress_bar.empty()
    
    # Sort results by similarity score
    identification_results.sort(key=lambda x: x["Persentase"], reverse=True)
    for i, result in enumerate(identification_results):
        result["Rank"] = i + 1
    
    log_container.info(f"🎯 Final results: {len(identification_results)} matches found")
    if identification_results:
        log_container.info(f"🏆 Top match: {identification_results[0]['Nama Bakteri']} ({identification_results[0]['Persentase']:.2f}%)")
        
    return identification_results

def fetch_and_display_detailed_profiles(session, genera_list):
    """Mengambil semua profil mentah, menampilkannya dalam tabel detail, dan mengembalikan tabel tersebut."""
    st.header("3. Data Detail dari BacDive")
    st.info("Tabel ini berisi data lengkap yang diambil dari BacDive untuk setiap strain, yang telah diratakan (flattened) dari format JSON aslinya.")

    all_dfs = []
    progress_bar = st.progress(0, text="Mengambil profil untuk semua genus...")

    for i, genus in enumerate(genera_list):
        status_placeholder = st.empty()
        # Log container and expander removed for a cleaner UI.
        raw_profiles = fetch_and_cache_profiles_by_taxonomy(session, genus, status_placeholder)

        if raw_profiles:
            all_profiles_list = []
            for bacdive_id, profile_data in raw_profiles.items():
                if isinstance(profile_data, dict):
                    # Buat salinan untuk dimodifikasi, ini lebih aman
                    row_data = profile_data.copy()
                    row_data['bacdive_id'] = bacdive_id
                    row_data['genus_input'] = genus
                    all_profiles_list.append(row_data)
            
            if all_profiles_list:
                all_dfs.append(pd.DataFrame(all_profiles_list))
        
        progress_bar.progress((i + 1) / len(genera_list), text=f"Selesai mengambil profil untuk {genus}")
        status_placeholder.empty()

    progress_bar.empty()
    
    if not all_dfs:
        st.warning("Tidak ada profil yang ditemukan untuk genus yang diberikan.")
        return pd.DataFrame()

    final_df = pd.concat(all_dfs, ignore_index=True)

    # Reorder columns to put important ones first
    cols = ['bacdive_id', 'genus_input'] + [col for col in final_df.columns if col not in ['bacdive_id', 'genus_input']]
    final_df = final_df[cols]

    with st.expander("Tampilkan/Sembunyikan Tabel Data Detail", expanded=True):
        st.dataframe(final_df)
        
        csv_buffer = io.StringIO()
        final_df.to_csv(csv_buffer, index=False)
        
        st.download_button(
            label="📥 Download Data Detail Lengkap (.csv)",
            data=csv_buffer.getvalue(),
            file_name="bacdive_detailed_data.csv",
            mime="text/csv",
            key="download-detailed-profiles"
        )
    return final_df

# --- 5. Tampilan Aplikasi (UI) ---
def highlight_comparison(s):
    """Fungsi untuk highlighting hasil perbandingan dalam tabel."""
    colors = []
    for v in s:
        if v == '❌':
            colors.append('background-color: #FFCDD2')  # Red for mismatch
        elif v == '➖':
            colors.append('background-color: #FFF9C4')  # Yellow for partial/variable
        elif v == '❓':
            colors.append('background-color: #F5F5F5')  # Gray for no data
        else:
            colors.append('')
    return colors

def main():
    st.title("🔬 Identifikasi Bakteri Berbasis Genus")
    st.info("Upload file CSV/Excel dengan kolom **Sample_Name** dan **Genus** untuk memulai identifikasi.")

    session = init_session()
    if not session:
        st.stop()

    with st.sidebar:
        st.header("Pengaturan & Bantuan")
        if st.button("🔄 Refresh Sesi & Token", help="Klik jika Anda mengalami error Unauthorized atau ingin memulai sesi baru."):
            st.cache_resource.clear()
            st.rerun()

        st.info(
            "1. **Download Template**: Unduh `template_input.csv`.\n"
            "2. **Isi Data**: Masukkan nama sampel, **genus**, dan hasil uji lab.\n"
            "3. **Upload File**: Unggah file yang sudah diisi.\n"
            "4. **Hasil**: Hasil akan muncul secara otomatis per sampel."
        )
        
        # PERBAIKAN: Template download dengan error handling
        try:
            if os.path.exists("template_input.csv"):
                with open("template_input.csv", "r") as f:
                    st.download_button(
                        label="📥 Download Template Input",
                        data=f.read(),
                        file_name="template_input.csv",
                        mime="text/csv"
                    )
        except Exception as e:
            st.warning(f"Template file tidak tersedia: {e}")

        st.header("Mode Akuakultur")
        mode = st.selectbox(
            "Pilih preset bobot untuk genus target:",
            ("Default", "Aeromonas Focus", "Streptococcus Focus", "Edwardsiella Focus")
        )

        # Update weights based on mode
        if mode == "Aeromonas Focus":
            WEIGHTS.update({'Oxidase': 4, 'Nitrate_reduction': 3, 'Glucose': 2, 'Gram_stain': 4})
        elif mode == "Streptococcus Focus":
            WEIGHTS.update({'Gram_stain': 4, 'Catalase': 4, 'Oxidase': 4, 'VP': 3})
        elif mode == "Edwardsiella Focus":
            WEIGHTS.update({'H2S_production': 4, 'Indole': 4, 'Motility': 3, 'Citrate': 3})

    

    # PERBAIKAN: Enhanced file upload section
    st.header("1. Upload File Input")
    uploaded_file = st.file_uploader(
        "Unggah file CSV/Excel", 
        type=["csv", "xlsx"],
        help="File harus berisi kolom 'Sample_Name' dan 'Genus' minimal"
    )

    if uploaded_file:
        try:
            # PERBAIKAN: Better file reading with encoding handling
            if uploaded_file.name.endswith('.csv'):
                data = pd.read_csv(uploaded_file, encoding='utf-8')
            else:
                data = pd.read_excel(uploaded_file)
            
            # PERBAIKAN: Validate required columns before normalization
            required_cols = ['Sample_Name', 'Genus']
            missing_cols = [col for col in required_cols if col not in data.columns]
            
            if missing_cols:
                st.error(f"Kolom yang diperlukan tidak ditemukan: {missing_cols}")
                st.info("File harus berisi minimal kolom: Sample_Name, Genus")
                st.stop()
            
            data = normalize_columns(data)
            
            st.header("2. Preview Data Input (Setelah Normalisasi)")
            st.dataframe(data)

            # PERBAIKAN: Enhanced genus validation
            unique_genera = data["Genus"].dropna().unique()
            empty_genus_count = data["Genus"].isna().sum()
            
            if empty_genus_count > 0:
                st.warning(f"⚠️ Ditemukan {empty_genus_count} baris dengan genus kosong. Baris ini akan dilewati.")

            if len(unique_genera) > 0:
                st.info(f"🔍 Ditemukan {len(unique_genera)} genus unik di file Anda: **{', '.join(unique_genera)}**.")
                
                # PERBAIKAN: Option to skip detailed profile fetch for large datasets
                if len(unique_genera) > 5:
                    st.warning(f"Dataset besar terdeteksi ({len(unique_genera)} genus). Proses mungkin memakan waktu lama.")
                    if not st.checkbox("Lanjutkan dengan fetch data detail", value=True):
                        st.info("Fetch data detail dilewati. Langsung ke identifikasi per sampel.")
                        detailed_df = pd.DataFrame()
                    else:
                        detailed_df = fetch_and_display_detailed_profiles(session, unique_genera)
                else:
                    detailed_df = fetch_and_display_detailed_profiles(session, unique_genera)

                st.header("4. Hasil Identifikasi per Sampel")

                # List untuk menyimpan hasil dari setiap sampel untuk laporan akhir
                all_sample_reports = []
                
                total_samples = len(data)
                sample_progress = st.progress(0, text="Memulai identifikasi sampel...")
                
                for index, row in data.iterrows():
                    st.divider()
                    sample_name = row.get("Sample_Name", f"Sampel #{index + 1}")
                    st.subheader(f"▶️ Hasil untuk Sampel: {sample_name}")
                    
                    progress_pct = (index + 1) / total_samples
                    sample_progress.progress(
                        progress_pct, 
                        text=f"Memproses sampel {index + 1}/{total_samples}: {sample_name}"
                    )
                    
                    user_input = row.to_dict()
                    
                    with st.expander(f"Lihat Log Detail Proses Fetch API untuk Sampel: {sample_name}"):
                        log_container = st.container()
                        results = process_sample(session, user_input, log_container)

                    # Simpan hasil (bahkan jika kosong) untuk laporan akhir
                    all_sample_reports.append({
                        "sample_name": sample_name,
                        "results": results
                    })

                    if results:
                        top_result = results[0]
                        st.success(f"**Identifikasi Utama:** `{top_result['Nama Bakteri']}` ({top_result['Persentase']:.2f}% kemiripan)")

                        with st.expander("Lihat Daftar Kandidat & Laporan Detail"):
                            st.subheader("Daftar Kandidat Teratas (Top 10)")
                            results_df = pd.DataFrame(results).head(10)[["Rank", "Nama Bakteri", "Persentase", "ID"]]
                            st.dataframe(results_df)

                            st.subheader("Laporan Perbandingan (vs Kandidat Utama)")
                            report_df = pd.DataFrame(top_result['details'])
                            st.dataframe(report_df.style.apply(highlight_comparison, subset=['Cocok']))
                    else:
                        st.warning(f"❌ Tidak ada hasil yang cocok ditemukan untuk sampel {sample_name}.")
                        st.info("Kemungkinan penyebab: Genus tidak ditemukan di database BacDive atau masalah koneksi API.")

                sample_progress.empty()
                st.success(f"✅ Selesai memproses {total_samples} sampel!")

                # --- BAGIAN 5: LAPORAN LENGKAP DALAM FORMAT DOCX ---
                st.divider()
                st.header("5. Laporan Lengkap (.docx)")
                st.info("Gunakan tombol di bawah ini untuk mengunduh laporan lengkap dalam format DOCX yang berisi ringkasan, daftar kandidat, dan detail perbandingan untuk setiap sampel.")

                try:
                    from docx import Document
                    from docx.shared import Pt
                    
                    def add_df_to_doc(document, df):
                        """Helper function to add a pandas DataFrame to a docx table."""
                        if df.empty:
                            document.add_paragraph("[Tidak ada data]", style='Italic')
                            return
                        table = document.add_table(rows=1, cols=df.shape[1], style='Table Grid')
                        for j, col_name in enumerate(df.columns):
                            table.cell(0, j).text = str(col_name)
                        for i, row in df.iterrows():
                            row_cells = table.add_row().cells
                            for j, cell_value in enumerate(row):
                                row_cells[j].text = str(cell_value)

                    document = Document()
                    document.add_heading('Laporan Lengkap Identifikasi Bakteri', 0)
                    document.add_paragraph(f"Laporan dibuat pada: {pd.to_datetime('today').strftime('%d %B %Y, %H:%M')}")

                    document.add_heading('Ringkasan Hasil Identifikasi', level=1)
                    summary_data = []
                    for report in all_sample_reports:
                        if report['results']:
                            top_res = report['results'][0]
                            summary_data.append({
                                'Nama Sampel': report['sample_name'],
                                'Kandidat Teratas': top_res['Nama Bakteri'],
                                'Skor Kemiripan': f"{top_res['Persentase']:.2f}%"
                            })
                        else:
                            summary_data.append({
                                'Nama Sampel': report['sample_name'],
                                'Kandidat Teratas': 'Tidak ditemukan',
                                'Skor Kemiripan': 'N/A'
                            })
                    summary_df = pd.DataFrame(summary_data)
                    add_df_to_doc(document, summary_df)

                    document.add_page_break()
                    document.add_heading('Detail Identifikasi per Sampel', level=1)

                    for report in all_sample_reports:
                        document.add_heading(f"Sampel: {report['sample_name']}", level=2)
                        
                        results = report['results']
                        if not results:
                            document.add_paragraph("Tidak ada hasil yang cocok ditemukan untuk sampel ini.")
                            document.add_paragraph('---')
                            continue

                        top_result = results[0]
                        p = document.add_paragraph()
                        p.add_run('Identifikasi Utama: ').bold = True
                        p.add_run(f"{top_result['Nama Bakteri']} ({top_result['Persentase']:.2f}%)")

                        document.add_heading('Daftar Kandidat Teratas', level=3)
                        kandidat_df = pd.DataFrame(results).head(10)[["Rank", "Nama Bakteri", "Persentase", "ID"]]
                        add_df_to_doc(document, kandidat_df)

                        document.add_heading('Laporan Perbandingan Detail (vs Kandidat Utama)', level=3)
                        detail_df = pd.DataFrame(top_result['details'])
                        add_df_to_doc(document, detail_df)
                        document.add_paragraph('') # Spacer

                    doc_io = io.BytesIO()
                    document.save(doc_io)
                    doc_io.seek(0)

                    st.download_button(
                        label="📥 Download Laporan Lengkap (.docx)",
                        data=doc_io,
                        file_name="laporan_identifikasi_lengkap.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download-docx-report"
                    )

                except ImportError:
                    st.error("Paket 'python-docx' tidak terinstal. Fitur unduh DOCX tidak dapat digunakan.")
                    st.code("pip install python-docx")
                except Exception as e:
                    st.error(f"Gagal membuat file DOCX: {e}")
                    st.exception(e)

        except Exception as e:
            st.error(f"Terjadi kesalahan saat memproses file: {e}")
            st.exception(e)
            st.info("Silakan periksa format file Anda dan pastikan berisi kolom yang diperlukan.")

if __name__ == "__main__":
    main()